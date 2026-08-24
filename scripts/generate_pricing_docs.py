"""
Generate per-service pricing documents from pricing/pricing.json.

Output (pricing/out/):
  LKR/  <Service>.pdf  and  <Service>.docx   - Sri Lankan clients
  USD/  <Service>.pdf                        - international clients

The JSON is the single source of truth. Change a number there, re-run this,
and every document updates together. Set "draft": false in the JSON once the
figures are confirmed - that drops the DRAFT banner from all output.

Usage:  python scripts/generate_pricing_docs.py
"""

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "pricing" / "pricing.json"
OUT = ROOT / "pricing" / "out"

# Brand palette, adjusted for print on white.
CRIMSON = colors.HexColor("#C41E28")
INK = colors.HexColor("#1A1A1A")
MUTED = colors.HexColor("#6B6B6B")
HAIRLINE = colors.HexColor("#E0E0E0")
WASH = colors.HexColor("#FDF3F3")

COMPANY = "DKayLABS"
EMAIL = "contact@dkaylabs.com"
PHONE = "+94 77 037 2960"
LOCATION = "Colombo, Sri Lanka"

DRAFT_NOTE = (
    "DRAFT - INDICATIVE PRICING ONLY. These figures are placeholders pending "
    "final review and must not be issued to a client as a quotation."
)
PRICING_NOTE = (
    "Every project is scoped and priced around what you actually need - features, "
    "complexity, and timeline all move the number. The tiers below are a starting "
    "point, not a fixed menu; tell us what you're building and we'll send a "
    "tailored quote."
)


def money(amount: int, currency: str) -> str:
    """1800000 -> 'LKR 1,800,000'"""
    return f"{currency} {amount:,}"


def safe_name(title: str) -> str:
    """'SEO & Digital Marketing' -> 'SEO-and-Digital-Marketing' (filesystem safe)."""
    cleaned = title.replace("&", "and").replace("/", "-")
    return re.sub(r"[^A-Za-z0-9]+", "-", cleaned).strip("-")


# --------------------------------------------------------------------------- PDF


def build_pdf(service: dict, currency: str, cfg: dict, path: Path) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"{service['title']} - Pricing ({currency})",
        author=COMPANY,
    )

    base = getSampleStyleSheet()
    s_brand = ParagraphStyle("brand", base["Normal"], fontName="Helvetica-Bold",
                             fontSize=17, textColor=CRIMSON, spaceAfter=1)
    s_kicker = ParagraphStyle("kicker", base["Normal"], fontName="Helvetica",
                              fontSize=8, textColor=MUTED, spaceAfter=10)
    s_title = ParagraphStyle("title", base["Normal"], fontName="Helvetica-Bold",
                             fontSize=23, textColor=INK, spaceBefore=6, spaceAfter=6)
    s_body = ParagraphStyle("body", base["Normal"], fontName="Helvetica",
                            fontSize=10, textColor=INK, leading=15, spaceAfter=10)
    s_muted = ParagraphStyle("muted", s_body, fontSize=9, textColor=MUTED)
    s_h2 = ParagraphStyle("h2", base["Normal"], fontName="Helvetica-Bold", fontSize=9,
                          textColor=CRIMSON, spaceBefore=14, spaceAfter=7)
    s_draft = ParagraphStyle("draft", base["Normal"], fontName="Helvetica-Bold",
                             fontSize=8, textColor=colors.white, alignment=TA_CENTER,
                             leading=12)
    s_cell = ParagraphStyle("cell", base["Normal"], fontName="Helvetica", fontSize=9,
                            textColor=INK, leading=13)
    s_tier = ParagraphStyle("tier", s_cell, fontName="Helvetica-Bold", fontSize=10)
    s_price = ParagraphStyle("price", s_cell, fontName="Helvetica-Bold", fontSize=12,
                             textColor=CRIMSON)
    s_foot = ParagraphStyle("foot", base["Normal"], fontName="Helvetica", fontSize=8,
                            textColor=MUTED, leading=12)

    flow = []

    if cfg["draft"]:
        banner = Table([[Paragraph(DRAFT_NOTE, s_draft)]], colWidths=[170 * mm])
        banner.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), CRIMSON),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        flow += [banner, Spacer(1, 12)]

    flow.append(Paragraph(COMPANY, s_brand))
    flow.append(Paragraph(
        f"Service Pricing &middot; {cfg['audience']} &middot; "
        f"Prepared {date.today():%d %B %Y}", s_kicker))
    flow.append(HRFlowable(width="100%", thickness=0.7, color=HAIRLINE,
                           spaceBefore=2, spaceAfter=2))

    flow.append(Paragraph(service["title"], s_title))
    flow.append(Paragraph(service["summary"], s_body))

    flow.append(Paragraph("PRICING", s_h2))

    rows = [[
        Paragraph("<b>Package</b>", s_cell),
        Paragraph("<b>What's included</b>", s_cell),
        Paragraph(f"<b>Price ({service['billing']})</b>", s_cell),
    ]]
    for tier, note, amount in zip(cfg["tiers"], service["tierNotes"], service[currency]):
        rows.append([
            Paragraph(tier, s_tier),
            Paragraph(note, s_cell),
            Paragraph(money(amount, currency), s_price),
        ])

    table = Table(rows, colWidths=[30 * mm, 90 * mm, 50 * mm], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), WASH),
        ("LINEBELOW", (0, 0), (-1, 0), 0.9, CRIMSON),
        ("GRID", (0, 0), (-1, -1), 0.4, HAIRLINE),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (2, 0), (2, -1), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    flow.append(table)

    flow.append(Paragraph("TYPICAL TIMELINE", s_h2))
    flow.append(Paragraph(service["timeline"], s_body))

    flow.append(Paragraph("HOW WE PRICE", s_h2))
    flow.append(Paragraph(PRICING_NOTE, s_body))

    tax = ("Prices are exclusive of applicable taxes."
           if currency == "LKR" else
           "Prices are exclusive of applicable taxes and any bank or transfer fees.")
    flow.append(Paragraph(
        f"{tax} Quotations are valid for 30 days from the date of issue. "
        "Payment terms and milestones are confirmed in the project agreement.", s_muted))

    flow.append(Spacer(1, 14))
    flow.append(HRFlowable(width="100%", thickness=0.7, color=HAIRLINE, spaceAfter=8))
    flow.append(Paragraph(
        f"{COMPANY} &middot; {LOCATION}<br/>{EMAIL} &middot; {PHONE}", s_foot))

    doc.build(flow)


# -------------------------------------------------------------------------- DOCX


def _shade(cell, hex_color: str) -> None:
    """Cell background - python-docx has no API for this, so drop to XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def build_docx(service: dict, currency: str, cfg: dict, path: Path) -> None:
    doc = Document()

    doc.core_properties.title = f"{service['title']} - Pricing ({currency})"
    doc.core_properties.author = COMPANY

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    crimson = RGBColor(0xC4, 0x1E, 0x28)
    muted = RGBColor(0x6B, 0x6B, 0x6B)

    if cfg["draft"]:
        warn = doc.add_paragraph()
        run = warn.add_run(DRAFT_NOTE)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = crimson

    head = doc.add_paragraph()
    brand = head.add_run(COMPANY)
    brand.bold = True
    brand.font.size = Pt(17)
    brand.font.color.rgb = crimson

    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Service Pricing · {cfg['audience']} · "
        f"Prepared {date.today():%d %B %Y}")
    meta_run.font.size = Pt(8.5)
    meta_run.font.color.rgb = muted

    doc.add_heading(service["title"], level=1)
    doc.add_paragraph(service["summary"])

    doc.add_heading("Pricing", level=2)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Package", "What's included", f"Price ({service['billing']})")
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _shade(cell, "FDF3F3")

    for tier, note, amount in zip(cfg["tiers"], service["tierNotes"], service[currency]):
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(tier).bold = True
        cells[1].text = note
        price = cells[2].paragraphs[0].add_run(money(amount, currency))
        price.bold = True
        price.font.color.rgb = crimson
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("Typical timeline", level=2)
    doc.add_paragraph(service["timeline"])

    doc.add_heading("How we price", level=2)
    doc.add_paragraph(PRICING_NOTE)

    tax = ("Prices are exclusive of applicable taxes."
           if currency == "LKR" else
           "Prices are exclusive of applicable taxes and any bank or transfer fees.")
    terms = doc.add_paragraph()
    terms_run = terms.add_run(
        f"{tax} Quotations are valid for 30 days from the date of issue. "
        "Payment terms and milestones are confirmed in the project agreement.")
    terms_run.font.size = Pt(9)
    terms_run.font.color.rgb = muted

    foot = doc.add_paragraph()
    foot_run = foot.add_run(
        f"{COMPANY} · {LOCATION}\n{EMAIL} · {PHONE}")
    foot_run.font.size = Pt(8.5)
    foot_run.font.color.rgb = muted

    doc.save(str(path))


# -------------------------------------------------------------------------- main


def main() -> None:
    data = json.loads(DATA.read_text(encoding="utf-8"))
    draft = data.get("draft", True)
    tiers = data["tiers"]

    written = 0
    for currency, meta in data["currencies"].items():
        target = OUT / currency
        target.mkdir(parents=True, exist_ok=True)

        cfg = {"draft": draft, "tiers": tiers, "audience": meta["audience"]}

        for service in data["services"]:
            stem = safe_name(service["title"])

            build_pdf(service, currency, cfg, target / f"{stem}.pdf")
            written += 1

            # Word versions are for the local team to edit; international
            # pricing goes out as a fixed PDF only.
            if currency == "LKR":
                build_docx(service, currency, cfg, target / f"{stem}.docx")
                written += 1

    print(f"Wrote {written} files to {OUT}")
    if draft:
        print('NOTE: draft mode is on. Set "draft": false in pricing.json '
              "once figures are confirmed.")


if __name__ == "__main__":
    main()
